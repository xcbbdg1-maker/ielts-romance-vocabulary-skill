#!/usr/bin/env python3
"""Build story-only reading DOCX files for a rendered page-count audit.

These files keep the bilingual story segments and remove all per-word cards.
They are internal measurement artifacts, not the primary study edition.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from build_series_docx import (
    add_chapter_cover,
    add_volume_cover,
    configure_section_furniture,
    load_index,
)
from build_chapter_docx import (
    COLORS,
    add_story_runs,
    configure_styles,
    enable_update_fields,
    set_paragraph_border,
    set_paragraph_shading,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
VOLUMES = [
    (1, 2, 6, "裂城身份"),
    (2, 7, 11, "失踪课表"),
    (3, 12, 16, "潮线证词"),
    (4, 17, 21, "数字迷城"),
    (5, 22, 26, "玻璃办公室"),
    (6, 27, 31, "白色走廊"),
    (7, 32, 36, "无声藏品"),
    (8, 37, 40, "舆论风暴"),
    (9, 41, 45, "法庭回声"),
    (10, 46, 50, "世界尽头的潮声"),
]


def add_story_only_body(doc: Document, chapter: dict) -> None:
    doc.add_heading("高密度双语正文", level=1)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            f"本章共 60 个小段、{chapter['expected_total']} 个学习单位；本计数版仅保留正文和行内中文括注，不含逐词精讲。"
        ),
        size=10.3,
        color=COLORS["muted"],
    )
    for segment in chapter["segments"]:
        doc.add_heading(f"第 {segment['number']:02d} 段", level=2)
        story_p = doc.add_paragraph(style="Story Block")
        set_paragraph_shading(story_p, COLORS["story_fill"])
        set_paragraph_border(story_p, side="left", color=COLORS["heading_blue"], size=14, space=7)
        add_story_runs(story_p, segment["story"])


def build_volume(chapters: list[dict], output: Path, volume_number: int, arc: str) -> None:
    doc = Document()
    configure_styles(doc)
    enable_update_fields(doc)
    title = f"《零点回声》正文计数版·卷{volume_number:02d}｜{arc}"
    configure_section_furniture(doc.sections[0], title)
    add_volume_cover(doc, {"title": title}, chapters)
    for index, chapter in enumerate(chapters):
        if index == 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            doc.add_section()
        add_chapter_cover(doc, chapter)
        add_story_only_body(doc, chapter)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--index", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    parser.add_argument("--chapter-01", type=Path)
    args = parser.parse_args()
    _, chapter_map = load_index(args.index.resolve())
    manifest = []
    if args.chapter_01:
        source = args.chapter_01.resolve()
        doc = Document(source)
        removed = 0
        for paragraph in list(doc.paragraphs):
            if paragraph.style.name in {"Vocabulary Label", "Definition", "Unit Divider"}:
                paragraph._element.getparent().remove(paragraph._element)
                removed += 1
        for paragraph in doc.paragraphs:
            if paragraph.style.name == "Heading 1":
                paragraph.text = "高密度双语正文"
            elif paragraph.style.name == "Normal" and "390 项词汇解释" in paragraph.text:
                paragraph.text = "本计数版保留 60 个小段及行内中文括注，去掉全部段后词汇精讲。"
        output = args.out_dir / "正文计数版_第一章.docx"
        doc.save(output)
        manifest.append({"volume": 0, "chapters": [1], "path": str(output.resolve()), "removed_paragraphs": removed})
    for volume, first, last, arc in VOLUMES:
        chapters = [chapter_map[number] for number in range(first, last + 1)]
        output = args.out_dir / f"正文计数版_卷{volume:02d}_第{first:02d}-{last:02d}章.docx"
        build_volume(chapters, output, volume, arc)
        manifest.append({"volume": volume, "chapters": list(range(first, last + 1)), "path": str(output.resolve())})
    (args.out_dir / "story_only_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"status": "pass", "files": len(manifest), "output": str(args.out_dir.resolve())}, ensure_ascii=False))


if __name__ == "__main__":
    main()
